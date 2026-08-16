from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "J_B_Huang_Data_Resume.docx"

# Named design override: one-page technical resume.
# The standard_business_brief palette is retained, while margins and type scale
# are tightened so the resume remains readable on one US Letter page.
NAVY = "0B2545"
BLUE = "2E74B5"
INK = "202124"
MUTED = "5F6368"
LIGHT = "D9E2F3"


def set_cell_margins(cell, top=60, start=100, bottom=60, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BLUE, size="10", space="1"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def font(run, size=9.1, color=INK, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.1)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.04

    for name, size, color, before, after in [
        ("Resume Section", 10.3, NAVY, 5, 2),
        ("Resume Role", 9.4, INK, 2, 1),
    ]:
        if name not in [s.name for s in doc.styles]:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Use Word's actual list style for bullets, with explicit resume geometry.
    bullets = doc.styles["List Bullet"]
    bullets.font.name = "Arial"
    bullets._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    bullets._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    bullets.font.size = Pt(8.75)
    bullets.font.color.rgb = RGBColor.from_string(INK)
    bullets.paragraph_format.left_indent = Inches(0.17)
    bullets.paragraph_format.first_line_indent = Inches(-0.12)
    bullets.paragraph_format.space_before = Pt(0)
    bullets.paragraph_format.space_after = Pt(1.2)
    bullets.paragraph_format.line_spacing = 1.02


def add_section(doc, label):
    p = doc.add_paragraph(style="Resume Section")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(label.upper())
    font(r, 10.3, NAVY, bold=True)
    set_paragraph_border(p, color=LIGHT, size="6", space="1")
    return p


def add_bullet(doc, lead, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    r = p.add_run(lead)
    font(r, 8.75, NAVY, bold=True)
    r = p.add_run(text)
    font(r, 8.75, INK)
    return p


def add_role(doc, title, org, date):
    p = doc.add_paragraph(style="Resume Role")
    p.paragraph_format.keep_with_next = True
    left = p.add_run(f"{title} | {org}")
    font(left, 9.4, INK, bold=True)
    right = p.add_run(f"    {date}")
    font(right, 9.1, MUTED)
    return p


def add_skill_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.1)
    r = p.add_run(f"{label}: ")
    font(r, 8.75, NAVY, bold=True)
    r = p.add_run(value)
    font(r, 8.75, INK)
    return p


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.50)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    # Header block.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("J. B. HUANG")
    font(r, 20, NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DATA ENGINEERING  |  DATA ANALYTICS  |  TIME-SERIES RESEARCH")
    font(r, 8.7, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[City, Country]  •  [email@example.com]  •  [+1 xxx xxx xxxx]  •  [LinkedIn]  •  [GitHub]")
    font(r, 8.3, MUTED)
    set_paragraph_border(p, color=BLUE, size="10", space="2")

    add_section(doc, "Profile")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    text = (
        "Graduate-level data engineering and analytics practitioner building reproducible, research-oriented data products. "
        "Experienced in designing layered data lakes, preserving source lineage, validating time-series quality, and translating "
        "energy-market data into statistically defensible research marts."
    )
    font(p.add_run(text), 8.95, INK)

    add_section(doc, "Technical Skills")
    add_skill_line(doc, "Data engineering", "Python, pandas, PyArrow, Parquet, JSON/JSONL, schema normalization, partitioned datasets, checksums, retryable ingestion")
    add_skill_line(doc, "Analytics", "pandas, statsmodels, fixed-effects panel regression, HAC errors, clustered standard errors, wild cluster bootstrap, feature engineering")
    add_skill_line(doc, "Data sources", "Open Power System Data, Open-Meteo Historical Weather API, Open-Meteo Single Runs API, API request/vintage metadata")
    add_skill_line(doc, "Workflow", "Bronze / Silver / Gold architecture, data-quality reports, reproducible batch IDs, research documentation, Git")

    add_section(doc, "Selected Project")
    add_role(doc, "Three-Layer Regional Energy Data Lake", "Independent Research Project", "2026–Present")
    add_bullet(doc, "Architecture. ", "Designed a reproducible Bronze–Silver–Gold pipeline for European power-system time series, retaining immutable raw files, source URLs, ingestion timestamps, SHA-256 checksums, and batch-level manifests.")
    add_bullet(doc, "Scale and quality. ", "Converted 50,401 wide source records into 13.7M typed, long-format Parquet rows partitioned by year and month; quality checks found 0 invalid timestamps, 0 negative physical MW values, and surfaced 8,006 profile-range anomalies for review.")
    add_bullet(doc, "Research mart. ", "Built a DE-LU hourly Gold mart with 17,546 observations and engineered load forecast error, renewable share, price spikes, and negative-price indicators; documented conditional associations rather than causal claims.")
    add_bullet(doc, "Multi-zone panel. ", "Standardized an unbalanced six-zone panel with 269,541 rows while preserving wind-source lineage; estimated region/hour/weekday/year fixed-effects models with region-clustered errors and a 999-rep wild cluster bootstrap.")
    add_bullet(doc, "Statistical judgment. ", "Showed how inference changes under few-cluster corrections: the renewable-share price association moved from cluster-robust p=0.0002 to bootstrap p=0.200, making small-sample fragility an explicit research result rather than a hidden caveat.")
    add_bullet(doc, "Forecast vintages. ", "Added a separate 2024 Open-Meteo ECMWF IFS forecast-vintage branch that records forecast_run_utc, valid_time_utc, forecast_horizon_hours, exact request URLs, and 480 forecast events for later forecast-error analysis.")

    add_section(doc, "Education")
    add_role(doc, "[Degree / Major]", "[University Name]", "[20XX–20XX]")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run("Relevant coursework: ")
    font(r, 8.75, NAVY, bold=True)
    r = p.add_run("[Data Engineering] · [Statistics / Econometrics] · [Machine Learning] · [Database Systems]")
    font(r, 8.75, INK)

    add_section(doc, "Additional Information")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Languages: ")
    font(r, 8.75, NAVY, bold=True)
    r = p.add_run("[English level] · [Chinese level]   ")
    font(r, 8.75, INK)
    r = p.add_run("Work authorization: ")
    font(r, 8.75, NAVY, bold=True)
    r = p.add_run("[Status]")
    font(r, 8.75, INK)

    # Quiet footer for the editable draft.
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("Resume draft • replace bracketed fields before sending")
    font(r, 7.2, MUTED, italic=True)

    doc.core_properties.title = "J. B. Huang — Data Engineering Resume"
    doc.core_properties.subject = "One-page technical resume draft"
    doc.core_properties.author = "J. B. Huang"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
